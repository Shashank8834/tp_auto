"""
Document generator for Annexure 1 - TP Report (TNMM).
Uses the prepared template and fills in client data + benchmarking data
while carefully preserving all formatting.
"""

import copy
import os
import io
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from models import GenerateReportRequest, ConnectedPerson, TestedPartyFinancials
from excel_parser import BenchmarkingData, SearchStrategyData

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template", "annexure1_template.docx")


def _format_aed(amount: float) -> str:
    """Format a number as AED currency string."""
    if amount >= 0:
        return f"AED {amount:,.2f}"
    return f"(AED {abs(amount):,.2f})"


def _format_aed_short(amount: float) -> str:
    """Format number for table cells."""
    return f"{amount:,.2f}"


def _replace_text_in_runs(paragraph, old_text: str, new_text: str):
    """
    Replace text across runs while preserving formatting.
    Handles cases where the placeholder might span multiple runs.
    """
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # Build character-to-run mapping
    runs = paragraph.runs
    if not runs:
        return False

    char_map = []  # (run_index, char_index_in_run)
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            char_map.append((ri, ci))

    # Find the occurrence in full text
    start = full_text.find(old_text)
    if start == -1:
        return False

    end = start + len(old_text)

    # Figure out which runs are affected
    if start < len(char_map) and end <= len(char_map):
        start_run, start_char = char_map[start]
        end_run, end_char = char_map[end - 1]

        if start_run == end_run:
            # Simple case: old text is within a single run
            run = runs[start_run]
            run.text = run.text[:start_char] + new_text + run.text[end_char + 1:]
        else:
            # Complex case: spans multiple runs
            # Put replacement in first run, clear text from others
            first_run = runs[start_run]
            first_run.text = first_run.text[:start_char] + new_text

            # Clear middle runs completely
            for ri in range(start_run + 1, end_run):
                runs[ri].text = ""

            # Trim last run
            last_run = runs[end_run]
            last_run.text = last_run.text[end_char + 1:]
    else:
        # Fallback: simple text replacement on first run
        runs[0].text = full_text.replace(old_text, new_text)
        for r in runs[1:]:
            r.text = ""

    return True


def _replace_placeholder_everywhere(doc: Document, placeholder: str, value: str):
    """Replace a placeholder in all paragraphs, tables, headers, and footers."""
    # Paragraphs
    for para in doc.paragraphs:
        while placeholder in para.text:
            _replace_text_in_runs(para, placeholder, value)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    while placeholder in para.text:
                        _replace_text_in_runs(para, placeholder, value)

    # Headers and footers
    for section in doc.sections:
        for hf in [section.header, section.footer,
                    section.first_page_header, section.first_page_footer]:
            if hf:
                for para in hf.paragraphs:
                    while placeholder in para.text:
                        _replace_text_in_runs(para, placeholder, value)
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                while placeholder in para.text:
                                    _replace_text_in_runs(para, placeholder, value)


def _copy_cell_format(source_cell, target_cell):
    """Copy formatting from source cell to target cell."""
    # Copy cell properties (borders, shading, etc.)
    source_tc = source_cell._tc
    target_tc = target_cell._tc

    # Copy tcPr (table cell properties)
    source_tcPr = source_tc.find(qn('w:tcPr'))
    if source_tcPr is not None:
        target_tcPr = target_tc.find(qn('w:tcPr'))
        if target_tcPr is not None:
            target_tc.remove(target_tcPr)
        new_tcPr = copy.deepcopy(source_tcPr)
        target_tc.insert(0, new_tcPr)


def _copy_paragraph_format(source_para, target_para):
    """Copy paragraph formatting from source to target."""
    # Copy paragraph properties
    source_pPr = source_para._p.find(qn('w:pPr'))
    if source_pPr is not None:
        target_pPr = target_para._p.find(qn('w:pPr'))
        if target_pPr is not None:
            target_para._p.remove(target_pPr)
        new_pPr = copy.deepcopy(source_pPr)
        target_para._p.insert(0, new_pPr)


def _copy_run_format(source_run, target_run):
    """Copy run formatting (font properties) from source to target."""
    source_rPr = source_run._r.find(qn('w:rPr'))
    if source_rPr is not None:
        target_rPr = target_run._r.find(qn('w:rPr'))
        if target_rPr is not None:
            target_run._r.remove(target_rPr)
        new_rPr = copy.deepcopy(source_rPr)
        target_run._r.insert(0, new_rPr)


def _set_cell_text_formatted(cell, text: str, source_cell=None):
    """Set cell text while preserving or copying formatting."""
    # Clear existing paragraphs except the first
    for p in cell.paragraphs[1:]:
        p._p.getparent().remove(p._p)

    para = cell.paragraphs[0]

    if source_cell:
        # Copy paragraph format from source
        source_para = source_cell.paragraphs[0]
        _copy_paragraph_format(source_para, para)

        # Clear existing runs
        for r in para.runs:
            r._r.getparent().remove(r._r)

        # Add new run with source formatting
        run = para.add_run(text)
        if source_para.runs:
            _copy_run_format(source_para.runs[0], run)
    else:
        # Just set text in existing formatting
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""
        else:
            run = para.add_run(text)


def _add_table_row_from_template(table, template_row_idx: int) -> object:
    """
    Add a new row to a table by deep-copying an existing row's XML.
    This preserves ALL formatting: borders, shading, cell widths, fonts, etc.
    """
    template_row = table.rows[template_row_idx]
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _build_connected_persons_table(doc: Document, table_idx: int,
                                   connected_persons: list[ConnectedPerson]):
    """
    Rebuild the connected persons table (Table 0) with dynamic rows.
    Copies formatting from existing data rows.
    """
    table = doc.tables[table_idx]

    # Table 0 has: Row 0 = header, Row 1-3 = data rows
    # Keep header row, remove existing data rows, add new ones

    # Save formatting from the first data row (row 1)
    if len(table.rows) < 2:
        return

    # Remove all data rows (keep header)
    while len(table.rows) > 2:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    # Use row index 1 as template for formatting
    template_row_idx = 1

    # Set first person in the existing template row
    if connected_persons:
        cp = connected_persons[0]
        row = table.rows[1]
        cells = row.cells
        _set_cell_text_formatted(cells[0], cp.name)
        _set_cell_text_formatted(cells[1], cp.designation)
        _set_cell_text_formatted(cells[2], _format_aed(cp.remuneration))
        _set_cell_text_formatted(cells[3], cp.roles)

    # Add additional rows
    for i, cp in enumerate(connected_persons[1:], start=1):
        new_row = _add_table_row_from_template(table, template_row_idx)
        cells = new_row.cells
        _set_cell_text_formatted(cells[0], cp.name, table.rows[template_row_idx].cells[0])
        _set_cell_text_formatted(cells[1], cp.designation, table.rows[template_row_idx].cells[1])
        _set_cell_text_formatted(cells[2], _format_aed(cp.remuneration), table.rows[template_row_idx].cells[2])
        _set_cell_text_formatted(cells[3], cp.roles, table.rows[template_row_idx].cells[3])


def _build_executive_summary_table(doc: Document, table_idx: int,
                                   data: BenchmarkingData,
                                   financials: TestedPartyFinancials,
                                   activity_description: str = ""):
    """
    Update the executive summary table (Table 1) with actual data.
    This table has a fixed structure, just update values.
    """
    table = doc.tables[table_idx]
    if len(table.rows) < 4:
        return

    # Row 2-3 contain the data cells
    # Update comparable count, margins, quartiles, and activity description
    for row in table.rows:
        for cell in row.cells:
            text = cell.text
            if "43" in text:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = run.text.replace("43", str(data.total_accepted))

            if "3.16%" in text:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = run.text.replace("3.16%", financials.op_or_percentage)

            # Replace activity-specific text
            if "Tour operating" in text and activity_description:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = run.text.replace("Tour operating services", activity_description)
                        run.text = run.text.replace("Tour operating service", activity_description)
                        run.text = run.text.replace("Tour operating", activity_description)

            if data.quartiles:
                if "2.65%" in text:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace("2.65%", data.quartiles.lower_quartile)
                if "5.43%" in text:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace("5.43%", data.quartiles.median)
                if "7.75%" in text:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace("7.75%", data.quartiles.upper_quartile)


def _build_search_strategy_table(doc: Document, table_idx: int,
                                 strategy: SearchStrategyData):
    """
    Fully rebuild a search strategy table with data from Excel.
    Template has 3 columns: col0 = criterion, col1 = description, col2 = count.
    Some criteria have merged col0+col1 in the template (no description).
    """
    table = doc.tables[table_idx]

    if len(table.rows) < 2:
        return

    # Row 0 = header ("Search Strategy"), Rows 1-12 = data
    # Keep header row (0) and first data row (1) as formatting template
    while len(table.rows) > 2:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    template_row_idx = 1

    def _fill_strategy_row(row, sr, template_cells=None):
        """Fill a search strategy row with criterion, description, count."""
        cells = row.cells
        # Put criterion text in col 0
        if template_cells:
            _set_cell_text_formatted(cells[0], sr.criterion, template_cells[0])
        else:
            _set_cell_text_formatted(cells[0], sr.criterion)

        # Put description in col 1 (if available, otherwise repeat criterion)
        if len(cells) > 1:
            desc = sr.description if sr.description else sr.criterion
            if template_cells:
                _set_cell_text_formatted(cells[1], desc, template_cells[1])
            else:
                _set_cell_text_formatted(cells[1], desc)

        # Count in last column
        if len(cells) > 2:
            if template_cells:
                _set_cell_text_formatted(cells[-1], sr.count, template_cells[-1])
            else:
                _set_cell_text_formatted(cells[-1], sr.count)

    # Set first search criterion in the existing template row
    if strategy.rows:
        _fill_strategy_row(table.rows[1], strategy.rows[0])

    # Add remaining search criteria rows
    for sr in strategy.rows[1:]:
        new_row = _add_table_row_from_template(table, template_row_idx)
        _fill_strategy_row(new_row, sr, table.rows[template_row_idx].cells)


def _build_rejection_table(doc: Document, table_idx: int,
                           data: BenchmarkingData):
    """
    Build the manual review rejection table (Table 6).
    Dynamic rows based on rejection reasons.
    """
    table = doc.tables[table_idx]

    if len(table.rows) < 3:
        return

    # Keep header row (row 0) and first data row as template
    # Remove all rows except header and first data row
    while len(table.rows) > 2:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    template_row_idx = 1

    # Set first rejection reason
    if data.rejection_reasons:
        rr = data.rejection_reasons[0]
        row = table.rows[1]
        cells = row.cells
        _set_cell_text_formatted(cells[0], str(rr.number))
        _set_cell_text_formatted(cells[1], rr.reason)
        _set_cell_text_formatted(cells[2], str(rr.count))

    # Add more rejection reason rows
    for rr in data.rejection_reasons[1:]:
        new_row = _add_table_row_from_template(table, template_row_idx)
        cells = new_row.cells
        _set_cell_text_formatted(cells[0], str(rr.number), table.rows[template_row_idx].cells[0])
        _set_cell_text_formatted(cells[1], rr.reason, table.rows[template_row_idx].cells[1])
        _set_cell_text_formatted(cells[2], str(rr.count), table.rows[template_row_idx].cells[2])

    # Add totals row
    totals_row = _add_table_row_from_template(table, template_row_idx)
    cells = totals_row.cells
    _set_cell_text_formatted(cells[0], "Total manual review rejections", table.rows[template_row_idx].cells[0])
    _set_cell_text_formatted(cells[1], "Total manual review rejections", table.rows[template_row_idx].cells[1])
    _set_cell_text_formatted(cells[2], str(data.total_rejections), table.rows[template_row_idx].cells[2])

    # Add accept row
    accept_row_num = len(data.rejection_reasons) + 1
    accept_row = _add_table_row_from_template(table, template_row_idx)
    cells = accept_row.cells
    _set_cell_text_formatted(cells[0], str(accept_row_num), table.rows[template_row_idx].cells[0])
    _set_cell_text_formatted(cells[1], "Accept", table.rows[template_row_idx].cells[1])
    _set_cell_text_formatted(cells[2], str(data.total_accepted), table.rows[template_row_idx].cells[2])

    # Add grand total row
    grand_total_row = _add_table_row_from_template(table, template_row_idx)
    cells = grand_total_row.cells
    _set_cell_text_formatted(cells[0], "Total", table.rows[template_row_idx].cells[0])
    _set_cell_text_formatted(cells[1], "Total", table.rows[template_row_idx].cells[1])
    _set_cell_text_formatted(cells[2], str(data.total_rejections + data.total_accepted),
                             table.rows[template_row_idx].cells[2])


def _build_comparable_companies_table(doc: Document, table_idx: int,
                                      data: BenchmarkingData):
    """
    Build the comparable companies table (Table 7) with dynamic rows.
    This is the large table with company margins across FY years.
    Format preservation is CRITICAL here.
    """
    table = doc.tables[table_idx]

    if len(table.rows) < 4:
        return

    # Rows 0-1 = headers (two-row header with merged cells)
    # Row 2 onwards = data rows
    # Last 3 rows = Lower quartile, Median, Upper quartile

    # Keep header rows (0, 1) and one data row (2) as template
    # Remove everything after row 2
    while len(table.rows) > 3:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    template_row_idx = 2  # First data row

    # Set first company data
    if data.comparable_companies:
        cc = data.comparable_companies[0]
        row = table.rows[2]
        cells = row.cells
        _set_cell_text_formatted(cells[0], str(cc.sno))
        _set_cell_text_formatted(cells[1], cc.name)

        # FY margin columns
        for i, fy in enumerate(data.fy_years[:3]):
            col_idx = 2 + i
            if col_idx < len(cells):
                val = cc.margins.get(fy, "N/A")
                _set_cell_text_formatted(cells[col_idx], val)

        # Weighted average column
        wa_idx = 2 + len(data.fy_years[:3])
        if wa_idx < len(cells):
            _set_cell_text_formatted(cells[wa_idx], cc.weighted_avg)

    # Add remaining company rows
    for cc in data.comparable_companies[1:]:
        new_row = _add_table_row_from_template(table, template_row_idx)
        cells = new_row.cells
        template_cells = table.rows[template_row_idx].cells

        _set_cell_text_formatted(cells[0], str(cc.sno), template_cells[0])
        _set_cell_text_formatted(cells[1], cc.name, template_cells[1])

        for i, fy in enumerate(data.fy_years[:3]):
            col_idx = 2 + i
            if col_idx < len(cells):
                val = cc.margins.get(fy, "N/A")
                _set_cell_text_formatted(cells[col_idx], val, template_cells[col_idx])

        wa_idx = 2 + len(data.fy_years[:3])
        if wa_idx < len(cells):
            _set_cell_text_formatted(cells[wa_idx], cc.weighted_avg, template_cells[wa_idx])

    # Add quartile rows
    if data.quartiles:
        quartile_items = [
            ("Lower quartile", data.quartiles.lower_quartile, data.quartiles.per_year_lower),
            ("Median", data.quartiles.median, data.quartiles.per_year_median),
            ("Upper quartile", data.quartiles.upper_quartile, data.quartiles.per_year_upper),
        ]

        for label, weighted_val, per_year in quartile_items:
            qrow = _add_table_row_from_template(table, template_row_idx)
            cells = qrow.cells
            template_cells = table.rows[template_row_idx].cells
            _set_cell_text_formatted(cells[0], label, template_cells[0])
            _set_cell_text_formatted(cells[1], label, template_cells[1])

            # Fill per-FY columns with actual values
            for i, fy in enumerate(data.fy_years[:3]):
                col_idx = 2 + i
                if col_idx < len(cells):
                    fy_val = per_year.get(fy, "")
                    _set_cell_text_formatted(cells[col_idx], fy_val, template_cells[col_idx])

            wa_idx = 2 + len(data.fy_years[:3])
            if wa_idx < len(cells):
                _set_cell_text_formatted(cells[wa_idx], weighted_val, template_cells[wa_idx])


def _build_financials_table(doc: Document, table_idx: int,
                            financials: TestedPartyFinancials,
                            fiscal_year_end: str):
    """
    Update the tested party financials table (Table 8).
    This table has a fixed structure, just update values.
    """
    table = doc.tables[table_idx]

    value_map = {
        "Operating Revenue": _format_aed_short(financials.operating_revenue),
        "Total Operating Revenue (OR)": _format_aed_short(financials.operating_revenue),
        "Operating Cost (OC)": _format_aed_short(financials.total_operating_cost),
        "Cost of sales": _format_aed_short(financials.cost_of_sales),
        "Admin & General expenses": _format_aed_short(financials.admin_expenses),
        "Other expenses": _format_aed_short(financials.other_expenses),
        "Staff Salary and benefits": _format_aed_short(financials.staff_salary),
        "Partners Salaries": _format_aed_short(financials.partner_salaries),
        "Operating Profit (OP = OR - OC)": _format_aed_short(financials.operating_profit),
        "(OP/OR )": financials.op_or_percentage,
    }

    for row in table.rows:
        cells = row.cells
        if not cells:
            continue

        label = cells[0].text.strip()

        # Update the "As at" date in header
        if "As at" in cells[-1].text:
            _set_cell_text_formatted(cells[-1], f"As at {fiscal_year_end}")

        # Update values
        for key, val in value_map.items():
            if label == key:
                _set_cell_text_formatted(cells[-1], val)
                break


def _generate_conclusion_text(financials: TestedPartyFinancials,
                              data: BenchmarkingData,
                              company_name: str) -> str:
    """Generate the conclusion paragraph text based on computed data."""
    margin = financials.op_or_percentage

    if data.quartiles:
        lq = data.quartiles.lower_quartile
        med = data.quartiles.median
        uq = data.quartiles.upper_quartile
    else:
        lq = "N/A"
        med = "N/A"
        uq = "N/A"

    conclusion = (
        f"The analysis confirms that the controlled transactions undertaken by the tested party "
        f"comply with the arm\u2019s length principle. This is evidenced by the tested party\u2019s margin "
        f"of {margin}, which falls significantly within the ALP of the operating profit to operating "
        f"revenue ratio observed among comparable companies. This strong performance, when compared "
        f"to the industry benchmark range (lower quartile: {lq}, median: {med}, upper quartile: "
        f"{uq}), clearly indicates that the company\u2019s remuneration and related-party payments "
        f"are consistent with arm\u2019s length standards, thereby substantiating the Company\u2019s "
        f"compliance with UAE transfer pricing regulations."
    )

    return conclusion


def generate_report(request: GenerateReportRequest,
                    benchmarking_data: BenchmarkingData) -> bytes:
    """
    Generate the Annexure 1 TP Report by filling the template with provided data.

    Returns the generated .docx file as bytes.
    """
    doc = Document(TEMPLATE_PATH)

    # ========== STEP 1: Replace all text placeholders ==========
    replacements = {
        "{{COMPANY_NAME}}": request.company_info.company_name,
        "{{COMPANY_SHORT_NAME}}": request.company_info.company_short_name or request.company_info.company_name,
        "{{FIRM_NAME}}": "BCL Globiz Accounting and Consulting L.L.C",
        "{{FIRM_SHORT_NAME}}": "BCL Globiz",
        "{{FIRM_INITIALS}}": "BCL",
        "{{FISCAL_YEAR_END_LONG}}": request.company_info.fiscal_year_end,
        "{{FISCAL_YEAR_END_SHORT}}": request.company_info.fiscal_year_end,
        "{{FISCAL_YEAR_START}}": request.company_info.fiscal_year_start,
        "{{FISCAL_YEAR_RANGE}}": f"{request.company_info.fiscal_year_start} to {request.company_info.fiscal_year_end}",
        "{{ACTIVITY_DESCRIPTION}}": request.company_info.activity_description or "Business Services",
        "{{ACTIVITY_DESCRIPTION_LOWER}}": (request.company_info.activity_description or "Business services").lower(),
        "{{TRANSACTION_DESCRIPTION}}": request.company_info.transaction_description or request.company_info.nature_of_business,
    }

    for placeholder, value in replacements.items():
        _replace_placeholder_everywhere(doc, placeholder, value)

    # ========== STEP 2: Update the About Company section ==========
    # Find and update the company description paragraph
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Replace company description (paragraph that starts with company description)
        if "comprehensive travel and tourism provider" in text.lower() or \
           "{{COMPANY_DESCRIPTION}}" in text:
            # Replace with user-provided description
            if request.company_info.nature_of_business:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = request.company_info.nature_of_business
                else:
                    para.add_run(request.company_info.nature_of_business)

        # Update total accepted count references
        if "43" in text and "companies" in text.lower():
            for run in para.runs:
                run.text = run.text.replace("43", str(benchmarking_data.total_accepted))

        if "3.16%" in text:
            for run in para.runs:
                run.text = run.text.replace("3.16%", request.financials.op_or_percentage)

        # Update quartile values in conclusion and other paragraphs
        if benchmarking_data.quartiles:
            if "2.65%" in text:
                for run in para.runs:
                    run.text = run.text.replace("2.65%", benchmarking_data.quartiles.lower_quartile)
            if "5.43%" in text:
                for run in para.runs:
                    run.text = run.text.replace("5.43%", benchmarking_data.quartiles.median)
            if "7.75%" in text:
                for run in para.runs:
                    run.text = run.text.replace("7.75%", benchmarking_data.quartiles.upper_quartile)

    # ========== STEP 3: Build dynamic tables ==========

    # Table 0: Connected Persons
    _build_connected_persons_table(doc, 0, request.connected_persons)

    # Table 1: Executive Summary
    _build_executive_summary_table(doc, 1, benchmarking_data, request.financials,
                                   request.company_info.activity_description)

    # Tables 3-5: Search Strategy tables (UAE, MENA, Europe)
    strategy_table_indices = [3, 4, 5]
    for idx, strategy in zip(strategy_table_indices, benchmarking_data.search_strategies):
        if idx < len(doc.tables):
            _build_search_strategy_table(doc, idx, strategy)

    # Table 6: Manual Review Rejections
    if len(doc.tables) > 6:
        _build_rejection_table(doc, 6, benchmarking_data)

    # Table 7: Comparable Companies
    if len(doc.tables) > 7:
        _build_comparable_companies_table(doc, 7, benchmarking_data)

    # Table 8: Tested Party Financials
    if len(doc.tables) > 8:
        _build_financials_table(doc, 8, request.financials, request.company_info.fiscal_year_end)

    # ========== STEP 4: Update manual review rejections total text ==========
    for para in doc.paragraphs:
        if "Manual review rejections total:" in para.text:
            # Replace the entire paragraph text, preserving run formatting
            if para.runs:
                para.runs[0].text = f"Manual review rejections total: {benchmarking_data.total_rejections}"
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.add_run(f"Manual review rejections total: {benchmarking_data.total_rejections}")

        if "Total Accepted after search and evaluation:" in para.text:
            if para.runs:
                para.runs[0].text = f"Total Accepted after search and evaluation: {benchmarking_data.total_accepted}"
                for r in para.runs[1:]:
                    r.text = ""

    # ========== STEP 5: Save to bytes ==========
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_report_to_file(request: GenerateReportRequest,
                            benchmarking_data: BenchmarkingData,
                            output_path: str):
    """Generate the report and save to a file."""
    doc_bytes = generate_report(request, benchmarking_data)
    with open(output_path, 'wb') as f:
        f.write(doc_bytes)
    return output_path
