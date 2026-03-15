"""
FastAPI backend for Annexure 1 - TP Report Generation.
Provides endpoints for form submission, Excel upload, document generation, and download.
"""

import os
import json
import uuid
import shutil
import pickle
import traceback
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

from models import GenerateReportRequest, CompanyInfo, ConnectedPerson, RelatedParty, TestedPartyFinancials
from excel_parser import parse_annexure2, BenchmarkingData
from doc_generator import generate_report_to_file
from annexure3_parser import parse_annexure3

app = FastAPI(
    title="TP Report Generator",
    description="Automated generation of Annexure 1 - Transfer Pricing Report (TNMM)",
    version="1.1.0"
)

# CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Storage directories
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")
SESSIONS_DIR = os.path.join(os.path.dirname(__file__), "sessions")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(SESSIONS_DIR, exist_ok=True)


# ============================================================
# Session persistence — survives server restarts
# ============================================================
def _save_session(session_id: str, session_data: dict):
    """Save session to disk as pickle."""
    path = os.path.join(SESSIONS_DIR, f"{session_id}.pkl")
    with open(path, "wb") as f:
        pickle.dump(session_data, f)


def _load_session(session_id: str) -> Optional[dict]:
    """Load session from disk."""
    path = os.path.join(SESSIONS_DIR, f"{session_id}.pkl")
    if os.path.exists(path):
        try:
            with open(path, "rb") as f:
                return pickle.load(f)
        except Exception:
            return None
    return None


def _delete_old_sessions(max_age_hours: int = 24):
    """Clean up sessions older than max_age_hours."""
    import time
    now = time.time()
    try:
        for f in os.listdir(SESSIONS_DIR):
            fp = os.path.join(SESSIONS_DIR, f)
            if os.path.isfile(fp) and now - os.path.getmtime(fp) > max_age_hours * 3600:
                os.remove(fp)
    except Exception:
        pass


# ============================================================
# Validation helpers
# ============================================================
def _validate_company_info(data: dict) -> list[str]:
    """Validate company info fields, return list of errors."""
    errors = []
    required = ["company_name", "nature_of_business", "address", "fiscal_year_start", "fiscal_year_end"]
    for field in required:
        val = data.get(field, "").strip()
        if not val:
            errors.append(f"Company info: '{field}' is required")
    return errors


def _validate_connected_persons(persons: list) -> list[str]:
    """Validate connected persons, return list of errors."""
    errors = []
    if not persons:
        errors.append("At least one connected person is required")
        return errors
    for i, p in enumerate(persons):
        if not p.get("name", "").strip():
            errors.append(f"Connected person #{i+1}: 'name' is required")
        if not p.get("designation", "").strip():
            errors.append(f"Connected person #{i+1}: 'designation' is required")
        rem = p.get("remuneration")
        if rem is None or (isinstance(rem, str) and not rem.strip()):
            errors.append(f"Connected person #{i+1}: 'remuneration' is required")
        elif isinstance(rem, str):
            try:
                float(rem)
            except ValueError:
                errors.append(f"Connected person #{i+1}: 'remuneration' must be a number")
    return errors


def _validate_financials(data: dict) -> list[str]:
    """Validate financial data, return list of errors."""
    errors = []
    required = ["operating_revenue", "cost_of_sales"]
    for field in required:
        val = data.get(field)
        if val is None or (isinstance(val, str) and not val.strip()):
            errors.append(f"Financials: '{field}' is required")
        elif isinstance(val, str):
            try:
                float(val)
            except ValueError:
                errors.append(f"Financials: '{field}' must be a number")
    return errors


# ============================================================
# API Endpoints
# ============================================================
@app.get("/api/health")
async def health():
    return {"status": "ok", "message": "TP Report Generator API is running"}


@app.post("/api/upload-excel")
async def upload_excel(file: UploadFile = File(...)):
    """Upload Annexure 2 Excel file for parsing."""
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Please upload an Excel file (.xlsx or .xls)")

    session_id = str(uuid.uuid4())
    session_dir = os.path.join(UPLOAD_DIR, session_id)
    os.makedirs(session_dir, exist_ok=True)

    # Save the uploaded file
    file_path = os.path.join(session_dir, file.filename)
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)

    # Parse the Excel file with graceful error handling
    try:
        benchmarking_data = parse_annexure2(file_path)
    except Exception as e:
        # Clean up
        shutil.rmtree(session_dir, ignore_errors=True)
        error_msg = str(e)
        if "no sheet" in error_msg.lower() or "worksheet" in error_msg.lower():
            raise HTTPException(
                status_code=400,
                detail=f"Excel file is missing required sheets. Expected sheets: "
                       f"'Summary of TP BM', 'UAE Search Strategy', 'Mena & Turkiye Search Strategy', "
                       f"'Europe Search Strategy', 'Accept reject matrix', 'Margin Analysis'. "
                       f"Error: {error_msg}"
            )
        raise HTTPException(
            status_code=400,
            detail=f"Could not parse Excel file. Please ensure it follows the Annexure 2 format. Error: {error_msg}"
        )

    # Validate parsed data
    warnings = []
    if not benchmarking_data.search_strategies:
        warnings.append("No search strategy data found — tables 3-5 will be empty")
    if not benchmarking_data.comparable_companies:
        warnings.append("No comparable companies found — table 7 will be empty")
    if not benchmarking_data.quartiles:
        warnings.append("No quartile data found — arm's length range will show N/A")

    # Store session to disk
    session_data = {
        "excel_path": file_path,
        "benchmarking_data": benchmarking_data,
    }
    _save_session(session_id, session_data)

    # Return summary for frontend display
    return {
        "session_id": session_id,
        "filename": file.filename,
        "warnings": warnings if warnings else None,
        "summary": {
            "search_strategies_count": len(benchmarking_data.search_strategies),
            "regions": [s.region for s in benchmarking_data.search_strategies],
            "total_comparables": benchmarking_data.total_accepted,
            "rejection_reasons_count": len(benchmarking_data.rejection_reasons),
            "total_rejections": benchmarking_data.total_rejections,
            "fy_years": benchmarking_data.fy_years,
            "quartiles": {
                "lower": benchmarking_data.quartiles.lower_quartile if benchmarking_data.quartiles else "N/A",
                "median": benchmarking_data.quartiles.median if benchmarking_data.quartiles else "N/A",
                "upper": benchmarking_data.quartiles.upper_quartile if benchmarking_data.quartiles else "N/A",
            } if benchmarking_data.quartiles else None,
            "comparable_companies": [
                {"sno": c.sno, "name": c.name, "weighted_avg": c.weighted_avg}
                for c in benchmarking_data.comparable_companies[:5]  # Preview first 5
            ],
        }
    }


@app.post("/api/upload-annexure3")
async def upload_annexure3(file: UploadFile = File(...)):
    """Upload Annexure 3 Excel file to extract financials and connected persons."""
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Please upload an Excel file (.xlsx or .xls)")

    # Save to a temp location
    temp_dir = os.path.join(UPLOAD_DIR, f"a3_{uuid.uuid4().hex[:8]}")
    os.makedirs(temp_dir, exist_ok=True)
    file_path = os.path.join(temp_dir, file.filename)

    try:
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)

        result = parse_annexure3(file_path)
        return {
            "filename": file.filename,
            "financials": result["financials"],
            "connected_persons": result["connected_persons"],
        }
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Could not parse Annexure 3 file: {str(e)}"
        )
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


@app.post("/api/generate")
async def generate_report_endpoint(
    form_data: str = Form(...),
    session_id: str = Form(...),
):
    """
    Generate the TP Report.
    form_data: JSON string with company info, connected persons, financials
    session_id: session ID from the upload-excel endpoint
    """
    # Parse form data
    try:
        data = json.loads(form_data)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid form data JSON")

    # ===== SERVER-SIDE VALIDATION =====
    errors = []
    errors.extend(_validate_company_info(data.get("company_info", {})))
    errors.extend(_validate_connected_persons(data.get("connected_persons", [])))
    errors.extend(_validate_financials(data.get("financials", {})))

    if errors:
        raise HTTPException(status_code=400, detail=f"Validation errors:\n" + "\n".join(f"• {e}" for e in errors))

    # Get benchmarking data from session (disk-based)
    session = _load_session(session_id)
    if not session:
        raise HTTPException(
            status_code=404,
            detail="Session not found or expired. Please re-upload the Excel file."
        )

    benchmarking_data = session["benchmarking_data"]

    # Build the request model
    try:
        company_info = CompanyInfo(**data.get("company_info", {}))
        connected_persons = [ConnectedPerson(**cp) for cp in data.get("connected_persons", [])]
        related_parties = [RelatedParty(**rp) for rp in data.get("related_parties", [])]
        financials = TestedPartyFinancials(**data.get("financials", {}))

        request = GenerateReportRequest(
            company_info=company_info,
            connected_persons=connected_persons,
            related_parties=related_parties,
            financials=financials,
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid form data: {str(e)}")

    # Generate the report
    try:
        report_id = str(uuid.uuid4())
        # Sanitize filename: remove special chars
        safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in company_info.company_name)
        output_filename = f"Annexure_1_TP_Report_{safe_name.replace(' ', '_')}_{report_id[:8]}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        generate_report_to_file(request, benchmarking_data, output_path)

        # Update session with report info
        session["report_path"] = output_path
        session["report_id"] = report_id
        session["output_filename"] = output_filename
        _save_session(session_id, session)

        return {
            "report_id": report_id,
            "filename": output_filename,
            "message": "Report generated successfully",
            "download_url": f"/api/download/{report_id}"
        }
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating report: {str(e)}")


@app.get("/api/download/{report_id}")
async def download_report(report_id: str):
    """Download the generated report."""
    # Search all session files for this report
    for fname in os.listdir(SESSIONS_DIR):
        if not fname.endswith(".pkl"):
            continue
        session_id = fname[:-4]
        session = _load_session(session_id)
        if session and session.get("report_id") == report_id:
            report_path = session.get("report_path")
            if report_path and os.path.exists(report_path):
                return FileResponse(
                    path=report_path,
                    filename=session.get("output_filename", "Annexure_1_TP_Report.docx"),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    raise HTTPException(status_code=404, detail="Report not found")


@app.post("/api/update-report")
async def update_report(
    report_id: str = Form(...),
    updates: str = Form(...),
):
    """Apply text edits to the generated report."""
    from docx import Document as DocxDocument

    # Find session
    session = None
    session_id = None
    for fname in os.listdir(SESSIONS_DIR):
        if not fname.endswith(".pkl"):
            continue
        sid = fname[:-4]
        s = _load_session(sid)
        if s and s.get("report_id") == report_id:
            session = s
            session_id = sid
            break

    if not session:
        raise HTTPException(status_code=404, detail="Report not found")

    report_path = session.get("report_path")
    if not report_path or not os.path.exists(report_path):
        raise HTTPException(status_code=404, detail="Report file not found")

    try:
        edits = json.loads(updates)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid updates JSON")

    try:
        doc = DocxDocument(report_path)

        for old_text, new_text in edits.items():
            for para in doc.paragraphs:
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if old_text in para.text:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)

        doc.save(report_path)
        return {"message": "Report updated successfully", "report_id": report_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating report: {str(e)}")


@app.get("/api/report-content/{report_id}")
async def get_report_content(report_id: str):
    """Get report content as structured text for the editor."""
    from docx import Document as DocxDocument

    session = None
    for fname in os.listdir(SESSIONS_DIR):
        if not fname.endswith(".pkl"):
            continue
        sid = fname[:-4]
        s = _load_session(sid)
        if s and s.get("report_id") == report_id:
            session = s
            break

    if not session:
        raise HTTPException(status_code=404, detail="Report not found")

    report_path = session.get("report_path")
    if not report_path or not os.path.exists(report_path):
        raise HTTPException(status_code=404, detail="Report file not found")

    doc = DocxDocument(report_path)

    sections = []
    current_section = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        if "Heading" in style_name:
            if current_section:
                sections.append(current_section)
            current_section = {
                "heading": text,
                "style": style_name,
                "paragraphs": []
            }
        else:
            if current_section is None:
                current_section = {
                    "heading": "Introduction",
                    "style": "Normal",
                    "paragraphs": []
                }
            current_section["paragraphs"].append({
                "text": text,
                "style": style_name,
            })

    if current_section:
        sections.append(current_section)

    tables_data = []
    for ti, table in enumerate(doc.tables):
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(cells)
        tables_data.append({"index": ti, "rows": table_rows})

    return {
        "report_id": report_id,
        "sections": sections,
        "tables": tables_data,
    }


# Clean up old sessions on startup
_delete_old_sessions()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
